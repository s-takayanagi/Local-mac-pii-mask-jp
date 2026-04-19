import shutil
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from docx import Document
from core.pipeline import mask_text
from file_handlers.base import masked_output_path
from models import ProcessResult

_LABEL_TO_TAG: dict[str, str] = {
    "氏名": "[氏名]",
    "ふりがな": "[氏名]",
}


def _merge_numeric(totals: dict, values: dict) -> None:
    for k, v in values.items():
        totals[k] = totals.get(k, 0) + v


def _mask_many(
    texts: list[str],
    model: str,
    url: str,
    enabled_layers: set[str] | None,
    excluded_tags: set[str] | None,
    max_workers: int,
):
    if max_workers > 1 and len(texts) > 1:
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            return list(ex.map(
                lambda t: mask_text(t, model, url, enabled_layers, excluded_tags),
                texts,
            ))
    return [mask_text(t, model, url, enabled_layers, excluded_tags) for t in texts]


def _apply_paragraph(para, final_text: str) -> None:
    """mask_text の final_text を段落の最初の run に書き、残りの run は空にする。"""
    if not para.runs:
        return
    para.runs[0].text = final_text
    for run in para.runs[1:]:
        run.text = ""


def process_docx(
    path: Path,
    model: str,
    lm_studio_url: str,
    enabled_layers: set[str] | None = None,
    excluded_tags: set[str] | None = None,
    max_workers: int = 1,
) -> ProcessResult:
    output = masked_output_path(path)
    shutil.copy2(path, output)

    doc = Document(output)
    total = 0
    errors: list[str] = []
    warnings: list[str] = []
    layer_totals: dict = {}
    layer_elapsed: dict = {}
    replacements_log: list[dict] = []

    # Phase 1: collect paragraph tasks + per-cell warning contexts
    # 各 task: (para_ref, loc, text, cell_key or None)
    tasks: list[tuple] = []
    # cell_key -> {"label": ..., "cell_text": ...}
    warning_contexts: dict[str, dict] = {}

    # document-level paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not para.runs or not text.strip():
            continue
        tasks.append((para, f"段落{para_idx + 1}", text, None))

    def _collect_table(table, loc_prefix: str) -> None:
        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            label_tag: str | None = None
            label_text: str = ""
            if cells:
                label_text = cells[0].text.strip()
                label_tag = _LABEL_TO_TAG.get(label_text)

            for col_idx, cell in enumerate(cells):
                loc = f"{loc_prefix}/行{row_idx + 1}/列{col_idx + 1}"
                for para in cell.paragraphs:
                    text = para.text
                    if not para.runs or not text.strip():
                        continue
                    cell_key = loc if (col_idx == 1 and label_tag is not None) else None
                    tasks.append((para, loc, text, cell_key))

                if col_idx == 1 and label_tag is not None and cell.text.strip():
                    warning_contexts[loc] = {
                        "label": label_text,
                        "cell_text": cell.text.strip(),
                    }

                for nested_table in cell.tables:
                    _collect_table(nested_table, f"{loc}(nested)")

    for tbl_idx, table in enumerate(doc.tables):
        _collect_table(table, f"表{tbl_idx + 1}")

    if not tasks:
        doc.save(output)
        return ProcessResult(
            output_path=output,
            total_replacements=0,
            errors=errors,
            layer_totals=layer_totals,
            layer_elapsed=layer_elapsed,
            replacements_log=replacements_log,
            warnings=warnings,
        )

    # Phase 2: mask
    texts = [t[2] for t in tasks]
    results = _mask_many(texts, model, lm_studio_url, enabled_layers, excluded_tags, max_workers)

    # Phase 3: apply results + aggregate
    cell_replaced: dict[str, int] = {}
    for (para, loc, original, cell_key), result in zip(tasks, results):
        try:
            count = len(result.replacements)
            total += count
            _merge_numeric(layer_totals, result.layer_counts)
            _merge_numeric(layer_elapsed, result.layer_elapsed)
            if result.error:
                errors.append(result.error)
            if result.final_text != original:
                _apply_paragraph(para, result.final_text)
                for r in result.replacements:
                    replacements_log.append({**r, "location": loc})
            if cell_key is not None:
                cell_replaced[cell_key] = cell_replaced.get(cell_key, 0) + count
        except Exception as e:
            errors.append(f"{loc}: {e}")

    # Label-cell warnings: 氏名/ふりがなラベル列の値セルが未マスクの場合
    for loc, ctx in warning_contexts.items():
        if cell_replaced.get(loc, 0) == 0:
            cell_text = ctx["cell_text"]
            warnings.append(
                f"{loc} の「{cell_text[:30]}」は{ctx['label']}の可能性があります。"
                "マスキングされていないことを確認してください。"
            )

    doc.save(output)
    return ProcessResult(
        output_path=output,
        total_replacements=total,
        errors=errors,
        layer_totals=layer_totals,
        layer_elapsed=layer_elapsed,
        replacements_log=replacements_log,
        warnings=warnings,
    )
